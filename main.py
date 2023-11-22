import pandas as pd
from docx import Document
import openai

# generator that yelds pairs of Number (int) and FaultName (str)
def xls_read_gen(file_path):
    data = pd.read_excel(file_path, skiprows=3)
    data = data.iloc[:, :2]
    for _, row in data.iterrows():
        yield row['No.'], row['Faults']

# usage example
if __name__ == "__main__":
    file_path = "faults.xlsx"
    for number, fault in xls_read_gen(file_path):
        print(f"Number: {number}, Fault: {fault}")



# generator that yelds pairs of Number (int) and FaultName (str)
def read_word_document_gen(file_path):
    color_red = "\033[91m"
    color_reset = "\033[0m"

    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("FLT") and not paragraph.text.startswith("FLT "):
            print(f"{color_red}Error in the document, missing space: {color_reset}{paragraph.text}")
    for paragraph in doc.paragraphs:
        line = paragraph.text
        for line in paragraph.text.split('\n'):
            if line.startswith("FLT "):
                yield int(line.split()[1][:-1]), ' '.join(line.split()[2:])

    # these lines for printing tables
    #  for table in doc.tables:
    #      for row in table.rows:
    #          for cell in row.cells:
    #              print(cell.text)


# Setze deinen OpenAI API-Schlüssel hier ein
openai.api_key = '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'

def compare_with_gpt(master_text, error_text):
    # Verwende die OpenAI GPT API für den Textvergleich
    response = openai.Completion.create(
        engine="text-davinci-002",
        prompt=f"Vergleiche den Text: \nMaster:\n{master_text}\nError:\n{error_text}\n",
        max_tokens=50
    )
    return response.choices[0].text.strip()

def write_to_document(output_file, numbers, descriptions):
    doc = docx.Document()
    for number, description in zip(numbers, descriptions):
        doc.add_paragraph(f"Number: {number}\nDescription: {description}\n\n")
    doc.save(output_file)
# usage example
if __name__ == "__main__":
    file_path = "masterdoc.docx"
    for number, line in read_word_document_gen(file_path):
        print(number, line)