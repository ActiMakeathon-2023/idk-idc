from docx import Document

# generator that yelds pairs of Number (int) and FaultName (str)
def read_word_document_gen(file_path):
    color_red = "\033[91m"
    color_reset = "\033[0m"

    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("FLT") and not paragraph.text.startswith("FLT "):
            print(f"{color_red}Error in the document, missing space: {color_reset}{paragraph.text}")
    for paragraph in doc.paragraphs:
        line = paragraph.text
        for line in paragraph.text.split('\n'):
            if line.startswith("FLT "):
                yield int(line.split()[1][:-1]), ' '.join(line.split()[2:])

    # these lines for printing tables
    #  for table in doc.tables:
    #      for row in table.rows:
    #          for cell in row.cells:
    #              print(cell.text)

# usage example
if __name__ == "__main__":
    file_path = "masterdoc.docx"
    for number, line in read_word_document_gen(file_path):
        print(number, line)
